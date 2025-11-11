from pathlib import Path
from docx import Document

base = Path('docs/reference')
for doc_path in base.glob('*.docx'):
    print('=' * 80)
    print(doc_path.name)
    try:
        doc = Document(doc_path)
    except Exception as exc:
        print(f'无法打开: {exc}')
        continue
    count = [0]
    def output(text, tag):
        text = text.strip()
        if not text:
            return
        print(f'[{count[0]:02d}] [{tag}] {text}')
        count[0] += 1
    for para in doc.paragraphs:
        style = para.style.name if para.style else ''
        output(para.text, f'段落:{style}')
        if count[0] >= 80:
            print('...截断...')
            break
    if count[0] < 80:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    output(cell.text, '表格')
                    if count[0] >= 80:
                        print('...截断...')
                        break
                if count[0] >= 80:
                    break
            if count[0] >= 80:
                break

from pathlib import Path
from docx import Document

base = Path('docs/reference')
for doc_path in base.glob('*.docx'):
    doc = Document(doc_path)
    texts = []
    for para in doc.paragraphs:
        texts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = '\t'.join(cell.text.strip() for cell in row.cells)
            texts.append(row_text)
    out_path = doc_path.with_suffix('.txt')
    out_path.write_text('\n'.join(texts), encoding='utf-8')
